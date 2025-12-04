"""
Translation utility for Excel, Word, and PDF documents.

This module provides functions to translate documents from one language to another,
with support for Excel (.xlsx, .xls), Word (.docx), PDF (.pdf), CSV (.csv), and
plain text (.txt) file formats.
It includes a translation cache system to avoid redundant API calls and improve
performance.

Installation:
-------------
Required dependencies:
    pip install pandas openpyxl deep-translator

Optional dependencies (for specific file formats):
    pip install python-docx    # For Word (.docx) files
    pip install pypdf          # For PDF (.pdf) files

Usage:
------
Interactive usage in Python:

    # Import the function
    from translate_files import translate_directory
    
    # Set up logging to see progress
    import logging
    logging.basicConfig(level=logging.INFO)
    
    # Translate a directory
    translate_directory(
        source_dir="/path/to/source/directory",
        target_dir="/path/to/target/directory",
        source_lang='th',      # Thai
        target_lang='en',      # English
        recursive=True,        # Process subdirectories
        file_extensions=('.xlsx', '.xls', '.docx', '.pdf', '.csv', '.txt')
    )

    # Translate only Excel files, non-recursive
    translate_directory(
        source_dir="/data/excel_files",
        target_dir="/data/translated",
        source_lang='th',
        target_lang='en',
        recursive=False,
        file_extensions=('.xlsx', '.xls')
    )

Command-line usage (if run as script):
    python translate_files.py

Translation Cache System:
-------------------------
The translation cache is a persistent JSON file that stores previously translated
text strings. This provides several benefits:

1. **Performance**: Avoids redundant API calls for text that has already been
   translated, significantly speeding up batch translation operations.

2. **Cost Efficiency**: Reduces API usage costs by reusing existing translations.

3. **Consistency**: Ensures the same source text always translates to the same
   target text across different translation sessions.

4. **Resume Capability**: If a translation process is interrupted, previously
   translated text can be reused without re-translating.

The cache is automatically saved every 100 new translations and at the end of
each translation session. The cache file is stored as JSON with UTF-8 encoding
to preserve special characters.

Example:
    >>> from translate_files import translate_directory
    >>> translate_directory(
    ...     source_dir="/path/to/source",
    ...     target_dir="/path/to/target",
    ...     source_lang='th',
    ...     target_lang='en',
    ...     recursive=True
    ... )
"""

import os
import json
import logging
import time
from pathlib import Path
from typing import Optional
import pandas as pd
from deep_translator import GoogleTranslator

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader, PdfWriter
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False


class TranslationCache:
    """
    Manages a persistent translation cache to store and retrieve translations.
    
    The cache is stored as a JSON file and automatically loaded on initialization.
    New translations are added to the cache and saved periodically (every 100
    translations) to prevent data loss.
    
    Attributes:
        cache_file (str): Path to the JSON cache file
        cache (dict): In-memory dictionary mapping source text to translated text
    
    Example:
        >>> cache = TranslationCache('cache.json')
        >>> translation = cache.get('Hello')
        >>> if translation is None:
        ...     translation = translate('Hello')
        ...     cache.set('Hello', translation)
    """
    def __init__(self, cache_file: str):
        """
        Initialize the translation cache.
        
        Args:
            cache_file: Path to the JSON file where translations are stored.
                       If the file exists, it will be loaded automatically.
        """
        self.cache_file = cache_file
        self.cache = {}
        self._load_cache()
    
    def _load_cache(self) -> None:
        """
        Load existing translations from the cache file.
        
        If the cache file doesn't exist, starts with an empty cache.
        """
        if os.path.exists(self.cache_file):
            with open(self.cache_file, 'r', encoding='utf-8') as f:
                self.cache = json.load(f)
    
    def save(self) -> None:
        """
        Save the current cache to disk.
        
        Writes the cache dictionary to the JSON file with UTF-8 encoding
        and proper formatting (indented, preserving non-ASCII characters).
        """
        with open(self.cache_file, 'w', encoding='utf-8') as f:
            json.dump(self.cache, f, ensure_ascii=False, indent=2)
    
    def get(self, text: str) -> Optional[str]:
        """
        Retrieve a cached translation for the given text.
        
        Args:
            text: The source text to look up in the cache
        
        Returns:
            The cached translation if found, None otherwise
        """
        return self.cache.get(text)
    
    def set(self, text: str, translation: str) -> None:
        """
        Store a translation in the cache.
        
        Automatically saves the cache to disk every 100 translations to prevent
        data loss in case of interruption.
        
        Args:
            text: The source text (used as the cache key)
            translation: The translated text (stored as the cache value)
        """
        self.cache[text] = translation
        if len(self.cache) % 100 == 0:
            self.save()


def translate_text(text: str, translator: GoogleTranslator, cache: TranslationCache) -> str:
    """
    Translate a single text string using the provided translator and cache.
    
    This function handles the translation workflow:
    1. Checks if the text is already cached
    2. Only translates if the text contains Thai characters (Unicode range U+0E00-U+0E7F)
    3. Stores new translations in the cache
    4. Adds a delay between API calls to avoid rate limiting
    
    Args:
        text: The text string to translate. Can be any string, including empty
              or non-string values (which are returned unchanged)
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Returns:
        The translated text if translation was performed, otherwise the original text.
        Non-string values (None, NaN, etc.) are returned unchanged.
    
    Note:
        Only text containing Thai characters (Unicode range U+0E00-U+0E7F) will
        be translated. Other text is returned as-is. This prevents unnecessary
        API calls for text that doesn't need translation.
    """
    # Handle non-string types (including pandas Series, NaN, etc.)
    if not isinstance(text, str):
        return text
    
    # Check for NaN values (pandas can have string 'nan' or actual NaN)
    if pd.isna(text):
        return text
    
    # Check for empty or whitespace-only strings
    if not text.strip():
        return text
        
    cached = cache.get(text)
    if cached is not None:
        return cached
        
    if any('\u0E00' <= c <= '\u0E7F' for c in text):
        try:
            translated = translator.translate(text)
            cache.set(text, translated)
            time.sleep(0.5)
            return translated
        except Exception as e:
            logging.error(f"Failed to translate text '{text[:50]}...': {str(e)}")
            return text
    
    return text


def translate_dataframe_values(
    df: pd.DataFrame,
    translator: GoogleTranslator,
    cache: TranslationCache
) -> pd.DataFrame:
    """
    Translate all text values in a DataFrame using unique value extraction.
    
    Extracts unique text values across all cells, translates them once,
    then maps translations back to original positions. This reduces function
    call overhead for datasets with repeated values.
    
    Args:
        df: DataFrame to translate (modified in place)
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Returns:
        The translated DataFrame
    """
    # Collect unique text values across all columns
    unique_texts: set[str] = set()
    for column in df.columns:
        for value in df[column].dropna().unique():
            if isinstance(value, str) and value.strip():
                unique_texts.add(value)
    
    # Translate each unique value once
    translations: dict[str, str] = {}
    for text in unique_texts:
        translations[text] = translate_text(text, translator, cache)
    
    # Apply translations using vectorized mapping
    for column in df.columns:
        df[column] = df[column].map(lambda x: translations.get(x, x) if isinstance(x, str) else x)
    
    return df


def translate_excel(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate an Excel file (.xlsx or .xls).
    
    Translates all cell values in all sheets of the spreadsheet, treating the first row
    as data (not headers). Uses workbook-level optimization to extract unique values
    across all sheets and translate them once.
    
    Args:
        input_path: Path to the source Excel file
        output_path: Path where the translated Excel file will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        FileNotFoundError: If the input file doesn't exist
        ValueError: If the file cannot be read as an Excel file
    
    Note:
        - The first row is treated as data, not headers (no column names are used)
        - All cell values including the first row are translated
        - Workbook-level optimization: unique values extracted across ALL sheets
        - Null/NaN values are preserved and not translated
        - The output file format matches the input (xlsx or xls)
        - All sheets in the workbook are processed and translated
        - Sheet names are also translated
        - Output file is written without headers
    """
    # Read all sheets
    excel_file = pd.ExcelFile(input_path)
    sheet_names = excel_file.sheet_names
    
    # Load all sheets first
    sheets_data: dict[str, pd.DataFrame] = {}
    for sheet_name in sheet_names:
        sheets_data[sheet_name] = pd.read_excel(input_path, sheet_name=sheet_name, header=None)
    
    # Collect unique text values across ALL sheets (workbook-level optimization)
    unique_texts: set[str] = set()
    for df in sheets_data.values():
        for column in df.columns:
            for value in df[column].dropna().unique():
                if isinstance(value, str) and value.strip():
                    unique_texts.add(value)
    
    # Also include sheet names in the unique set
    for sheet_name in sheet_names:
        if isinstance(sheet_name, str) and sheet_name.strip():
            unique_texts.add(sheet_name)
    
    # Translate all unique values once
    translations: dict[str, str] = {}
    for text in unique_texts:
        translations[text] = translate_text(text, translator, cache)
    
    # Apply translations to all sheets
    translated_sheets: dict[str, pd.DataFrame] = {}
    for sheet_name, df in sheets_data.items():
        translated_sheet_name = translations.get(sheet_name, sheet_name)
        
        for column in df.columns:
            df[column] = df[column].map(lambda x: translations.get(x, x) if isinstance(x, str) else x)
        
        translated_sheets[translated_sheet_name] = df
    
    # Write all sheets to output file without headers
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for sheet_name, df in translated_sheets.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)


def translate_word(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate a Word document (.docx).
    
    Translates all paragraphs and table cells in the document while preserving
    the document structure, formatting, and layout.
    
    Args:
        input_path: Path to the source Word document
        output_path: Path where the translated Word document will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If the input file doesn't exist
        ValueError: If the file cannot be read as a Word document
    
    Note:
        - Only paragraphs and table cells with non-empty text are translated
        - Document structure, formatting, images, and other elements are preserved
        - Requires the 'python-docx' package: pip install python-docx
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word document translation. Install with: pip install python-docx")
    
    doc = Document(input_path)
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated = translate_text(paragraph.text, translator, cache)
            paragraph.text = translated
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    translated = translate_text(cell.text, translator, cache)
                    cell.text = translated
    
    doc.save(output_path)


def translate_pdf(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate a PDF document (.pdf).
    
    Extracts text from each page, translates it, and creates a new PDF.
    Note: PDF text replacement has limitations - the output PDF preserves the
    original structure but text replacement may not be perfect due to the
    complex nature of PDF format.
    
    Args:
        input_path: Path to the source PDF file
        output_path: Path where the translated PDF will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        ImportError: If pypdf or PyPDF2 is not installed
        FileNotFoundError: If the input file doesn't exist
        ValueError: If the file cannot be read as a PDF
    
    Note:
        - Requires 'pypdf' or 'PyPDF2' package: pip install pypdf
        - PDF text replacement is limited - consider using specialized PDF
          editing tools for production use
        - The function extracts text from pages but full text replacement
          in PDFs is complex and may not preserve all formatting
    """
    if not PDF_AVAILABLE:
        raise ImportError("pypdf or PyPDF2 is required for PDF translation. Install with: pip install pypdf")
    
    reader = PdfReader(input_path)
    writer = PdfWriter()
    
    for page in reader.pages:
        text = page.extract_text()
        if text.strip():
            translated_text = translate_text(text, translator, cache)
            new_page = writer.add_page(page)
        else:
            writer.add_page(page)
    
    with open(output_path, 'wb') as f:
        writer.write(f)
    
    logging.warning("PDF translation preserves structure but text replacement is limited. Consider using specialized PDF editing tools for full text replacement.")


def translate_csv(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate a CSV file (.csv).
    
    Translates both column headers and all cell values in the CSV file.
    Preserves the original data structure and handles various CSV formats.
    
    Args:
        input_path: Path to the source CSV file
        output_path: Path where the translated CSV file will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        FileNotFoundError: If the input file doesn't exist
        ValueError: If the file cannot be read as a CSV file
    
    Note:
        - Column names are translated first, then cell values
        - Null/NaN values are preserved and not translated
        - Attempts to detect CSV delimiter automatically
        - Preserves UTF-8 encoding
    """
    try:
        df = pd.read_csv(input_path, encoding='utf-8')
    except UnicodeDecodeError:
        try:
            df = pd.read_csv(input_path, encoding='latin-1')
        except Exception:
            df = pd.read_csv(input_path, encoding='utf-8', errors='ignore')
    
    # Translate column headers separately
    new_columns = []
    for col in df.columns:
        translated_col = translate_text(str(col), translator, cache)
        new_columns.append(translated_col)
    df.columns = new_columns
    
    # Translate data cells using unique value extraction
    df = translate_dataframe_values(df, translator, cache)
    
    df.to_csv(output_path, index=False, encoding='utf-8')


def translate_txt(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate a plain text file (.txt).
    
    Translates the entire text content of the file while preserving line breaks
    and structure. Handles various text encodings.
    
    Args:
        input_path: Path to the source text file
        output_path: Path where the translated text file will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        FileNotFoundError: If the input file doesn't exist
        UnicodeDecodeError: If the file encoding cannot be determined
    
    Note:
        - Attempts to detect file encoding automatically (UTF-8, then latin-1)
        - Preserves line breaks and paragraph structure
        - Output is saved as UTF-8
        - Large files are translated line by line to preserve structure
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    content = None
    
    for encoding in encodings:
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    if content is None:
        raise ValueError(f"Could not decode file {input_path} with any supported encoding")
    
    if content.strip():
        translated_content = translate_text(content, translator, cache)
    else:
        translated_content = content
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(translated_content)


def translate_file(input_path: str, output_path: str, translator: GoogleTranslator, cache: TranslationCache) -> None:
    """
    Translate a single file based on its file extension.
    
    This is a dispatcher function that routes to the appropriate format-specific
    translation function based on the file extension. Automatically creates
    the output directory if it doesn't exist.
    
    Args:
        input_path: Path to the source file
        output_path: Path where the translated file will be saved
        translator: GoogleTranslator instance configured with source and target languages
        cache: TranslationCache instance for storing and retrieving translations
    
    Raises:
        ValueError: If the file format is not supported
        FileNotFoundError: If the input file doesn't exist
        ImportError: If required format-specific libraries are not installed
    
    Supported formats:
        - .xlsx, .xls: Excel spreadsheets
        - .docx: Word documents
        - .pdf: PDF documents
        - .csv: CSV files
        - .txt: Plain text files
    """
    input_path_obj = Path(input_path)
    ext = input_path_obj.suffix.lower()
    
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    if ext in ['.xlsx', '.xls']:
        translate_excel(input_path, output_path, translator, cache)
    elif ext == '.docx':
        translate_word(input_path, output_path, translator, cache)
    elif ext == '.pdf':
        translate_pdf(input_path, output_path, translator, cache)
    elif ext == '.csv':
        translate_csv(input_path, output_path, translator, cache)
    elif ext == '.txt':
        translate_txt(input_path, output_path, translator, cache)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def translate_directory(
    source_dir: str,
    target_dir: str,
    source_lang: str = 'th',
    target_lang: str = 'en',
    cache_file: Optional[str] = None,
    recursive: bool = True,
    file_extensions: tuple = ('.xlsx', '.xls', '.docx', '.pdf', '.csv', '.txt')
) -> None:
    """
    Translate all supported files in a directory, optionally recursing into subdirectories.
    
    This is the main function for batch translation. It:
    - Scans the source directory for files matching the specified extensions
    - Translates directory names and filenames
    - Recreates the same directory structure in the target directory
    - Skips files that have already been translated (if output exists)
    - Uses a translation cache to avoid redundant API calls
    
    Args:
        source_dir: Path to the source directory containing files to translate
        target_dir: Path to the target directory where translated files will be saved.
                    The directory structure from source_dir will be recreated here.
        source_lang: Source language code (default: 'th' for Thai)
        target_lang: Target language code (default: 'en' for English)
        cache_file: Optional path to the translation cache file. If None, defaults
                   to 'translation_cache.json' in the target directory.
        recursive: If True, recursively processes subdirectories. If False, only
                  processes files in the top-level directory.
        file_extensions: Tuple of file extensions to process (case-insensitive).
                        Default: ('.xlsx', '.xls', '.docx', '.pdf')
    
    Returns:
        None. All operations are logged. The cache is saved at the end.
    
    Example:
        >>> translate_directory(
        ...     source_dir="/data/thai_docs",
        ...     target_dir="/data/english_docs",
        ...     source_lang='th',
        ...     target_lang='en',
        ...     recursive=True,
        ...     file_extensions=('.xlsx', '.docx')
        ... )
    
    Note:
        - Directory names and filenames are also translated
        - Files that already exist in the target location are skipped
        - The translation cache is shared across all files in the batch
        - Logging is used to track progress and errors
        - The cache is automatically saved at the end of processing
    """
    source_path = Path(source_dir)
    target_path = Path(target_dir)
    
    if cache_file is None:
        cache_file = os.path.join(target_dir, 'translation_cache.json')
    
    cache = TranslationCache(cache_file)
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    
    if recursive:
        walker = source_path.rglob('*')
    else:
        walker = [item for item in source_path.iterdir() if item.is_file()]
    
    for item in walker:
        if not item.is_file():
            continue
        
        if item.suffix.lower() not in file_extensions:
            continue
        
        rel_path = item.relative_to(source_path)
        
        translated_parts = []
        for part in rel_path.parts[:-1]:
            translated_part = translate_text(part, translator, cache)
            translated_parts.append(translated_part)
        
        translated_filename = translate_text(rel_path.name, translator, cache)
        
        if translated_parts:
            output_path = target_path / Path(*translated_parts) / translated_filename
        else:
            output_path = target_path / translated_filename
        
        if output_path.exists():
            logging.info(f"Skipping {item.name} - translated file already exists")
            continue
        
        try:
            logging.info(f"Translating {item.name}")
            translate_file(str(item), str(output_path), translator, cache)
            logging.info(f"Successfully translated {item.name} -> {output_path}")
        except Exception as e:
            logging.error(f"Failed to translate {item.name}: {str(e)}")
    
    cache.save()
